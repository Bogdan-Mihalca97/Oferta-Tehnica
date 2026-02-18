"""
PV Proposal Generator - GUI App
Homescreen with section buttons. Each section generates a part of the final document.
"""
import os
import sys
import json
import threading
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
import fitz  # PyMuPDF
import anthropic
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.config import anthropic_api_key


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROMPTS_DIR = os.path.join(BASE_DIR, 'prompts')


def _load_prompt(filename):
    """Load a prompt from the prompts/ directory."""
    path = os.path.join(PROMPTS_DIR, filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


# ---------------------------------------------------------------------------
# PDF EXTRACTION
# ---------------------------------------------------------------------------
def extract_pdf_text(pdf_path):
    """Extract all text from a PDF using PyMuPDF."""
    doc = fitz.open(pdf_path)
    pages = []
    for i in range(len(doc)):
        pages.append(doc[i].get_text())
    doc.close()
    return pages


# ---------------------------------------------------------------------------
# AI GENERATION
# ---------------------------------------------------------------------------
SYSTEM_PROMPT_PTE = _load_prompt('system_pte.txt')
SYSTEM_PROMPT_REZUMAT = _load_prompt('system_rezumat.txt')


def _stream_claude(client, model, system, user_prompt, progress_callback=None, chunk_label="", max_tokens=16384):
    """Make a single streaming Claude API call and return the result text."""
    result_parts = []
    chars_received = 0

    with client.messages.stream(
        model=model,
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user_prompt}],
    ) as stream:
        for event in stream:
            if hasattr(event, 'type'):
                if event.type == 'content_block_delta' and hasattr(event.delta, 'text'):
                    result_parts.append(event.delta.text)
                    chars_received += len(event.delta.text)
                    if chars_received % 500 < 50 and progress_callback:
                        progress_callback(f"  {chunk_label}Se generează... {chars_received} caractere primite")

        final_message = stream.get_final_message()
        input_tokens = final_message.usage.input_tokens
        output_tokens = final_message.usage.output_tokens

    if progress_callback:
        progress_callback(
            f"  {chunk_label}Terminat. "
            f"Input: {input_tokens} tokeni, Output: {output_tokens} tokeni"
        )

    return ''.join(result_parts), input_tokens, output_tokens


def generate_pte(methodology_pages, api_key, model, progress_callback=None):
    """Call Claude API to transform methodology into PTE format.

    Splits into 2 chunks only for large documents (>30K chars) to avoid output truncation.
    """
    client = anthropic.Anthropic(api_key=api_key)

    full_text = '\n'.join(methodology_pages)
    total_chars = len(full_text)

    # Decide: 1 call for short docs, 2 calls for long ones
    if total_chars <= 30000:
        chunks = [full_text]
        if progress_callback:
            progress_callback(f"Document scurt ({total_chars} caractere) - un singur apel API...")
    else:
        mid = len(methodology_pages) // 2
        chunks = [
            '\n'.join(methodology_pages[:mid]),
            '\n'.join(methodology_pages[mid:]),
        ]
        if progress_callback:
            progress_callback(f"Document mare ({total_chars} caractere) - se trimite în 2 părți ({mid} + {len(methodology_pages) - mid} pagini)...")

    all_results = []
    total_input = 0
    total_output = 0
    num_chunks = len(chunks)

    for i, chunk_text in enumerate(chunks):
        chunk_num = i + 1
        chunk_label = f"[Partea {chunk_num}/{num_chunks}] " if num_chunks > 1 else ""
        if progress_callback:
            if num_chunks > 1:
                progress_callback(f"Partea {chunk_num}/{num_chunks}: Se trimite către Claude API...")
            else:
                progress_callback("Se trimite către Claude API...")

        part_info = f" (partea {chunk_num}/{num_chunks})" if num_chunks > 1 else ""
        user_prompt = _load_prompt('user_pte.txt').format(
            part_info=part_info,
            chunk_text=chunk_text
        )

        result, inp_tok, out_tok = _stream_claude(
            client, model, SYSTEM_PROMPT_PTE, user_prompt,
            progress_callback=progress_callback,
            chunk_label=chunk_label
        )
        all_results.append(result)
        total_input += inp_tok
        total_output += out_tok

    if progress_callback:
        progress_callback(
            f"Generat cu succes! Total: {total_input} input tokeni, {total_output} output tokeni"
        )

    return '\n\n'.join(all_results)


def _load_reference_style():
    """Load reference style from s01_context.json if available."""
    ctx_path = os.path.join(BASE_DIR, 's01_context.json')
    if os.path.exists(ctx_path):
        with open(ctx_path, 'r', encoding='utf-8') as f:
            ctx = json.load(f)
        return ctx.get('reference_style', '')
    return ''


def generate_rezumat(notice_pages, datasheet_pages, atr_pages, company_data,
                     api_key, model, progress_callback=None):
    """Call Claude API to generate the Rezumat (Summary) section.

    Single API call - output is ~5 pages, no chunking needed.
    """
    client = anthropic.Anthropic(api_key=api_key)

    notice_text = '\n'.join(notice_pages)
    datasheet_text = '\n'.join(datasheet_pages)
    atr_text = '\n'.join(atr_pages)

    # Load reference style for tone/structure matching
    reference_style = _load_reference_style()
    reference_block = ""
    if reference_style:
        reference_block = f"""

EXEMPLU DE STIL (din documentul de referință - folosește EXACT acest stil, structură și nivel de detaliu, dar cu datele din proiectul curent):
{reference_style}
"""

    user_prompt = _load_prompt('user_rezumat.txt').format(
        warranty_months=company_data['warranty_months'],
        pm_experience=company_data['pm_experience'],
        leader=company_data['leader'],
        associate=company_data['associate'],
        subcontractor=company_data['subcontractor'],
        notice_text=notice_text,
        datasheet_text=datasheet_text,
        atr_text=atr_text,
        reference_block=reference_block
    )

    if progress_callback:
        progress_callback("Se trimite către Claude API...")

    result, inp_tok, out_tok = _stream_claude(
        client, model, SYSTEM_PROMPT_REZUMAT, user_prompt,
        progress_callback=progress_callback,
        max_tokens=16384
    )

    if progress_callback:
        progress_callback(
            f"Generat cu succes! Input: {inp_tok} tokeni, Output: {out_tok} tokeni"
        )

    return result


# ---------------------------------------------------------------------------
# DOCX BUILDER
# ---------------------------------------------------------------------------
def build_docx(pte_text, output_path, doc_type="pte"):
    """Build a DOCX document from generated text.

    Args:
        doc_type: "pte" = flat list with single title (no headings),
                  "generic" = full heading support for future doc types.
    """
    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Arial Narrow'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # PTE mode: add the single title at the top
    if doc_type == "pte":
        h = doc.add_heading("Proceduri tehnice de execuție în cadrul prezentului Contract", level=2)
        for run in h.runs:
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(14)

    # Parse markdown-like output
    lines = pte_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Skip headings in PTE mode (safety net)
        if line.startswith('#'):
            if doc_type == "pte":
                i += 1
                continue  # Skip all headings for PTE
            # Generic mode: render headings
            if line.startswith('#### '):
                heading_text = line[5:].strip()
                h = doc.add_heading(heading_text, level=3)
                for run in h.runs:
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(12)
            elif line.startswith('### '):
                heading_text = line[4:].strip()
                h = doc.add_heading(heading_text, level=2)
                for run in h.runs:
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(14)
            elif line.startswith('## '):
                heading_text = line[3:].strip()
                h = doc.add_heading(heading_text, level=1)
                for run in h.runs:
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(16)
            i += 1
            continue

        # Table detection: consecutive lines starting with |
        if line.strip().startswith('|') and doc_type != "pte":
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_rows.append(lines[i].strip())
                i += 1
            if table_rows:
                _add_table(doc, table_rows)
            continue

        if line.strip().startswith('- ') or line.strip().startswith('• '):
            # Bullet point
            bullet_text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, bullet_text)

        else:
            # Regular paragraph - handle **bold** markers
            p = doc.add_paragraph()
            _add_formatted_text(p, line.strip())

        i += 1

    doc.save(output_path)


def _add_formatted_text(paragraph, text):
    """Parse **bold** markers and add formatted runs to paragraph."""
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(12)
        if i % 2 == 1:  # Odd indices are bold
            run.bold = True


def _add_table(doc, rows):
    """Add a table to the document from pipe-separated markdown rows."""
    parsed_rows = []
    for row in rows:
        cells = [c.strip() for c in row.strip('|').split('|')]
        parsed_rows.append(cells)

    # Skip separator rows (like |---|---|)
    data_rows = [r for r in parsed_rows if not all(set(c) <= {'-', ' ', ':'} for c in r)]

    if not data_rows:
        return

    num_cols = max(len(r) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial Narrow'
                        run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# GUI - PAGE FRAMEWORK
# ---------------------------------------------------------------------------
class App:
    """Main application with page-based navigation."""

    def __init__(self, root):
        self.root = root
        self.root.title("Generator Propunere Tehnică - Parc Fotovoltaic")
        self.root.geometry("850x700")
        self.root.resizable(True, True)

        # Shared state
        self.api_key = tk.StringVar(value=anthropic_api_key)
        self.model = tk.StringVar(value='claude-sonnet-4-20250514')

        # Company data (shared across sections)
        self.company_leader = tk.StringVar(value='CRC AG S.R.L.')
        self.company_associate = tk.StringVar(value='CRC NEW ENERGY S.R.L.')
        self.company_subcontractor = tk.StringVar(value='BACKUP TECHNOLOGY S.R.L.')
        self.warranty_months = tk.IntVar(value=120)
        self.pm_experience = tk.IntVar(value=5)

        # Container for pages
        self.container = ttk.Frame(self.root)
        self.container.pack(fill=tk.BOTH, expand=True)

        # Pages dict
        self.pages = {}
        self._create_pages()
        self.show_page("home")

    def _create_pages(self):
        """Create all pages (frames) and store them."""
        for PageClass in (HomePage, PTEPage, RezumatPage):
            page = PageClass(self.container, self)
            self.pages[page.name] = page
            page.frame.grid(row=0, column=0, sticky="nsew")

        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)

    def show_page(self, page_name):
        """Raise the given page to the front."""
        page = self.pages[page_name]
        page.frame.tkraise()


# ---------------------------------------------------------------------------
# GUI - HOME PAGE
# ---------------------------------------------------------------------------
class HomePage:
    name = "home"

    def __init__(self, parent, app):
        self.app = app
        self.frame = ttk.Frame(parent, padding=30)

        # Title
        ttk.Label(
            self.frame,
            text="Generator Propunere Tehnică",
            font=('Arial', 18, 'bold')
        ).pack(pady=(20, 5))

        ttk.Label(
            self.frame,
            text="Parc Fotovoltaic",
            font=('Arial', 14)
        ).pack(pady=(0, 30))

        # Sections container
        sections_frame = ttk.LabelFrame(self.frame, text="Secțiuni disponibile", padding=20)
        sections_frame.pack(fill=tk.BOTH, expand=True, padx=20)

        # --- Section buttons ---
        # Each button is a row with icon-like label + description + arrow

        self._add_section_button(
            sections_frame,
            title="3.5.2  Generator PTE",
            description="Proceduri Tehnice de Execuție - generare din Metodologia de Execuție",
            command=lambda: app.show_page("pte"),
            enabled=True
        )

        self._add_section_button(
            sections_frame,
            title="2.  Metodologia de executare",
            description="Descrierea lucrărilor și echipamentelor necesare",
            command=None,
            enabled=False
        )

        self._add_section_button(
            sections_frame,
            title="1.  Rezumat",
            description="Date generale, obiectul contractului, avantaje competitive",
            command=lambda: app.show_page("rezumat"),
            enabled=True
        )

        self._add_section_button(
            sections_frame,
            title="3.6  Puncte de control calitate",
            description="Verificări și teste pe faze de execuție",
            command=None,
            enabled=False
        )

        self._add_section_button(
            sections_frame,
            title="5.  Personal propus",
            description="Echipa de proiect și responsabilități",
            command=None,
            enabled=False
        )

        # Settings at the bottom
        settings_frame = ttk.LabelFrame(self.frame, text="Setări globale", padding=10)
        settings_frame.pack(fill=tk.X, padx=20, pady=(15, 0))

        key_row = ttk.Frame(settings_frame)
        key_row.pack(fill=tk.X, pady=2)
        ttk.Label(key_row, text="Cheie API:", width=15).pack(side=tk.LEFT)
        ttk.Entry(key_row, textvariable=app.api_key, show="*", width=50).pack(side=tk.LEFT, fill=tk.X, expand=True)

        model_row = ttk.Frame(settings_frame)
        model_row.pack(fill=tk.X, pady=2)
        ttk.Label(model_row, text="Model:", width=15).pack(side=tk.LEFT)
        ttk.Combobox(
            model_row, textvariable=app.model, state='readonly',
            values=['claude-haiku-4-5-20251001', 'claude-sonnet-4-20250514', 'claude-opus-4-20250514'],
            width=40
        ).pack(side=tk.LEFT)

        # Company data
        company_frame = ttk.LabelFrame(self.frame, text="Date companie", padding=10)
        company_frame.pack(fill=tk.X, padx=20, pady=(10, 0))

        for label_text, var in [
            ("Lider asociere:", app.company_leader),
            ("Asociat:", app.company_associate),
            ("Subcontractant:", app.company_subcontractor),
        ]:
            row = ttk.Frame(company_frame)
            row.pack(fill=tk.X, pady=1)
            ttk.Label(row, text=label_text, width=15).pack(side=tk.LEFT)
            ttk.Entry(row, textvariable=var, width=40).pack(side=tk.LEFT, fill=tk.X, expand=True)

        numbers_row = ttk.Frame(company_frame)
        numbers_row.pack(fill=tk.X, pady=1)
        ttk.Label(numbers_row, text="Garanție (luni):", width=15).pack(side=tk.LEFT)
        ttk.Spinbox(numbers_row, textvariable=app.warranty_months, from_=12, to=240, width=6).pack(side=tk.LEFT)
        ttk.Label(numbers_row, text="   Exp. MP (proiecte):", width=20).pack(side=tk.LEFT)
        ttk.Spinbox(numbers_row, textvariable=app.pm_experience, from_=1, to=50, width=6).pack(side=tk.LEFT)

    def _add_section_button(self, parent, title, description, command, enabled=True):
        """Add a section button row."""
        btn_frame = ttk.Frame(parent)
        btn_frame.pack(fill=tk.X, pady=4)

        if enabled:
            btn = ttk.Button(btn_frame, text=f"  {title}  ", command=command, width=40)
            btn.pack(side=tk.LEFT, padx=(0, 10))
        else:
            btn = ttk.Button(btn_frame, text=f"  {title}  ", state=tk.DISABLED, width=40)
            btn.pack(side=tk.LEFT, padx=(0, 10))

        lbl = ttk.Label(btn_frame, text=description, foreground='gray' if not enabled else 'black')
        lbl.pack(side=tk.LEFT, fill=tk.X, expand=True)

        if not enabled:
            tag = ttk.Label(btn_frame, text="(în curând)", foreground='gray', font=('Arial', 8, 'italic'))
            tag.pack(side=tk.RIGHT)


# ---------------------------------------------------------------------------
# GUI - PTE PAGE
# ---------------------------------------------------------------------------
class PTEPage:
    name = "pte"

    def __init__(self, parent, app):
        self.app = app
        self.frame = ttk.Frame(parent, padding=15)
        self.methodology_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.generating = False

        self._build_ui()

    def _build_ui(self):
        # Top bar: back button + title
        top = ttk.Frame(self.frame)
        top.pack(fill=tk.X, pady=(0, 10))

        ttk.Button(top, text="< Înapoi", command=lambda: self.app.show_page("home")).pack(side=tk.LEFT)
        ttk.Label(
            top,
            text="3.5.2  Generator PTE - Proceduri Tehnice de Execuție",
            font=('Arial', 13, 'bold')
        ).pack(side=tk.LEFT, padx=15)

        # --- Input file ---
        input_frame = ttk.LabelFrame(self.frame, text="Fișier de intrare (Metodologie PDF)", padding=10)
        input_frame.pack(fill=tk.X, pady=5)

        input_row = ttk.Frame(input_frame)
        input_row.pack(fill=tk.X)

        ttk.Entry(input_row, textvariable=self.methodology_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(input_row, text="Alege fișier...", command=self._browse_input).pack(side=tk.RIGHT, padx=(10, 0))

        # --- Output file ---
        output_frame = ttk.LabelFrame(self.frame, text="Fișier de ieșire (DOCX)", padding=10)
        output_frame.pack(fill=tk.X, pady=5)

        output_row = ttk.Frame(output_frame)
        output_row.pack(fill=tk.X)

        ttk.Entry(output_row, textvariable=self.output_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_row, text="Alege locația...", command=self._browse_output).pack(side=tk.RIGHT, padx=(10, 0))

        # --- Generate button ---
        self.gen_btn = ttk.Button(self.frame, text="Generează PTE", command=self._start_generation)
        self.gen_btn.pack(pady=15)

        # --- Progress ---
        self.progress = ttk.Progressbar(self.frame, mode='indeterminate')
        self.progress.pack(fill=tk.X, pady=5)

        # --- Log ---
        log_frame = ttk.LabelFrame(self.frame, text="Jurnal", padding=5)
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.log = scrolledtext.ScrolledText(log_frame, height=12, font=('Consolas', 9), state=tk.DISABLED)
        self.log.pack(fill=tk.BOTH, expand=True)

    def _browse_input(self):
        input_dir = os.path.join(BASE_DIR, 'input')
        if not os.path.isdir(input_dir):
            input_dir = BASE_DIR
        path = filedialog.askopenfilename(
            title="Alege Metodologia PDF",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
            initialdir=input_dir
        )
        if path:
            self.methodology_path.set(path)
            output_dir = os.path.join(BASE_DIR, 'output')
            os.makedirs(output_dir, exist_ok=True)
            base_name = os.path.splitext(os.path.basename(path))[0]
            self.output_path.set(
                os.path.join(output_dir, f'PTE_{base_name}.docx')
            )

    def _browse_output(self):
        current = self.output_path.get()
        if current and os.path.isdir(os.path.dirname(current)):
            initial_dir = os.path.dirname(current)
            initial_file = os.path.basename(current)
        else:
            initial_dir = os.path.expanduser('~\\Desktop')
            initial_file = ''
        path = filedialog.asksaveasfilename(
            title="Salvează PTE ca DOCX",
            filetypes=[("Word Document", "*.docx")],
            defaultextension=".docx",
            initialdir=initial_dir,
            initialfile=initial_file
        )
        if path:
            self.output_path.set(path)

    def _log(self, message):
        """Thread-safe logging."""
        def _update():
            self.log.config(state=tk.NORMAL)
            self.log.insert(tk.END, message + '\n')
            self.log.see(tk.END)
            self.log.config(state=tk.DISABLED)
        self.app.root.after(0, _update)

    def _start_generation(self):
        if self.generating:
            return

        if not self.methodology_path.get():
            messagebox.showerror("Eroare", "Alege un fișier PDF cu metodologia!")
            return
        if not self.output_path.get():
            messagebox.showerror("Eroare", "Alege locația fișierului de ieșire!")
            return
        if not self.app.api_key.get():
            messagebox.showerror("Eroare", "Introdu cheia API Anthropic în setări!")
            return

        self.generating = True
        self.gen_btn.config(state=tk.DISABLED)
        self.progress.start(10)

        thread = threading.Thread(target=self._generate, daemon=True)
        thread.start()

    def _generate(self):
        try:
            # Step 1: Extract PDF text
            self._log("Pas 1/3: Se extrage textul din PDF...")
            pdf_path = self.methodology_path.get()
            pages = extract_pdf_text(pdf_path)
            total_chars = sum(len(p) for p in pages)
            self._log(f"  Extras {len(pages)} pagini, {total_chars} caractere")

            # Step 2: Generate PTE via Claude (split into 2 chunks)
            self._log("Pas 2/3: Se generează PTE prin Claude API...")
            self._log(f"  Model: {self.app.model.get()}")

            pte_text = generate_pte(
                pages,
                api_key=self.app.api_key.get(),
                model=self.app.model.get(),
                progress_callback=self._log
            )

            # Save raw text for reference
            raw_path = self.output_path.get().replace('.docx', '_raw.txt')
            with open(raw_path, 'w', encoding='utf-8') as f:
                f.write(pte_text)
            self._log(f"  Text brut salvat: {raw_path}")

            # Step 3: Build DOCX
            self._log("Pas 3/3: Se construiește documentul DOCX...")
            output_path = self.output_path.get()
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            build_docx(pte_text, output_path)
            self._log(f"  Document salvat: {output_path}")

            self._log("\nGata! Documentul a fost generat cu succes.")
            self.app.root.after(0, lambda: messagebox.showinfo(
                "Succes", f"PTE generat cu succes!\n\n{output_path}"))

        except anthropic.AuthenticationError:
            self._log("EROARE: Cheie API invalidă!")
            self.app.root.after(0, lambda: messagebox.showerror(
                "Eroare API", "Cheia API Anthropic este invalidă."))
        except anthropic.BadRequestError as e:
            self._log(f"EROARE API: {e}")
            self.app.root.after(0, lambda: messagebox.showerror("Eroare API", str(e)))
        except Exception as e:
            self._log(f"EROARE: {e}")
            self.app.root.after(0, lambda: messagebox.showerror("Eroare", str(e)))
        finally:
            self.generating = False
            self.app.root.after(0, lambda: self.gen_btn.config(state=tk.NORMAL))
            self.app.root.after(0, self.progress.stop)


# ---------------------------------------------------------------------------
# GUI - REZUMAT PAGE
# ---------------------------------------------------------------------------
class RezumatPage:
    name = "rezumat"

    def __init__(self, parent, app):
        self.app = app
        self.frame = ttk.Frame(parent, padding=15)
        self.notice_path = tk.StringVar()
        self.datasheet_path = tk.StringVar()
        self.atr_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.generating = False

        self._build_ui()

    def _build_ui(self):
        # Top bar: back button + title
        top = ttk.Frame(self.frame)
        top.pack(fill=tk.X, pady=(0, 10))

        ttk.Button(top, text="< Înapoi", command=lambda: self.app.show_page("home")).pack(side=tk.LEFT)
        ttk.Label(
            top,
            text="1. Rezumat - Date generale și obiectul contractului",
            font=('Arial', 13, 'bold')
        ).pack(side=tk.LEFT, padx=15)

        # --- Input files ---
        input_frame = ttk.LabelFrame(self.frame, text="Fișiere de intrare (PDF)", padding=10)
        input_frame.pack(fill=tk.X, pady=5)

        self._add_file_row(input_frame, "Anunț de participare:", self.notice_path, "Alege Anunțul PDF")
        self._add_file_row(input_frame, "Fișa de date:", self.datasheet_path, "Alege Fișa de date PDF")
        self._add_file_row(input_frame, "ATR (Aviz Tehnic Racordare):", self.atr_path, "Alege ATR PDF")

        # --- Output file ---
        output_frame = ttk.LabelFrame(self.frame, text="Fișier de ieșire (DOCX)", padding=10)
        output_frame.pack(fill=tk.X, pady=5)

        output_row = ttk.Frame(output_frame)
        output_row.pack(fill=tk.X)

        ttk.Entry(output_row, textvariable=self.output_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_row, text="Alege locația...", command=self._browse_output).pack(side=tk.RIGHT, padx=(10, 0))

        # --- Generate button ---
        self.gen_btn = ttk.Button(self.frame, text="Generează Rezumat", command=self._start_generation)
        self.gen_btn.pack(pady=15)

        # --- Progress ---
        self.progress = ttk.Progressbar(self.frame, mode='indeterminate')
        self.progress.pack(fill=tk.X, pady=5)

        # --- Log ---
        log_frame = ttk.LabelFrame(self.frame, text="Jurnal", padding=5)
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.log = scrolledtext.ScrolledText(log_frame, height=12, font=('Consolas', 9), state=tk.DISABLED)
        self.log.pack(fill=tk.BOTH, expand=True)

    def _add_file_row(self, parent, label_text, var, dialog_title):
        """Add a labeled file selector row."""
        row = ttk.Frame(parent)
        row.pack(fill=tk.X, pady=2)

        ttk.Label(row, text=label_text, width=28).pack(side=tk.LEFT)
        ttk.Entry(row, textvariable=var, width=45).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(
            row, text="...",
            command=lambda: self._browse_pdf(var, dialog_title),
            width=3
        ).pack(side=tk.RIGHT, padx=(5, 0))

    def _browse_pdf(self, var, title):
        input_dir = os.path.join(BASE_DIR, 'input')
        if not os.path.isdir(input_dir):
            input_dir = BASE_DIR
        path = filedialog.askopenfilename(
            title=title,
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
            initialdir=input_dir
        )
        if path:
            var.set(path)
            # Auto-set output path if not set
            if not self.output_path.get():
                output_dir = os.path.join(BASE_DIR, 'output')
                os.makedirs(output_dir, exist_ok=True)
                self.output_path.set(os.path.join(output_dir, 'S01_Rezumat.docx'))

    def _browse_output(self):
        current = self.output_path.get()
        if current and os.path.isdir(os.path.dirname(current)):
            initial_dir = os.path.dirname(current)
            initial_file = os.path.basename(current)
        else:
            initial_dir = os.path.expanduser('~\\Desktop')
            initial_file = 'S01_Rezumat.docx'
        path = filedialog.asksaveasfilename(
            title="Salvează Rezumat ca DOCX",
            filetypes=[("Word Document", "*.docx")],
            defaultextension=".docx",
            initialdir=initial_dir,
            initialfile=initial_file
        )
        if path:
            self.output_path.set(path)

    def _log(self, message):
        """Thread-safe logging."""
        def _update():
            self.log.config(state=tk.NORMAL)
            self.log.insert(tk.END, message + '\n')
            self.log.see(tk.END)
            self.log.config(state=tk.DISABLED)
        self.app.root.after(0, _update)

    def _start_generation(self):
        if self.generating:
            return

        # Validate inputs
        if not self.notice_path.get():
            messagebox.showerror("Eroare", "Alege fișierul PDF cu Anunțul de participare!")
            return
        if not self.datasheet_path.get():
            messagebox.showerror("Eroare", "Alege fișierul PDF cu Fișa de date!")
            return
        if not self.atr_path.get():
            messagebox.showerror("Eroare", "Alege fișierul PDF cu ATR!")
            return
        if not self.output_path.get():
            messagebox.showerror("Eroare", "Alege locația fișierului de ieșire!")
            return
        if not self.app.api_key.get():
            messagebox.showerror("Eroare", "Introdu cheia API Anthropic în setări (pagina principală)!")
            return

        self.generating = True
        self.gen_btn.config(state=tk.DISABLED)
        self.progress.start(10)

        thread = threading.Thread(target=self._generate, daemon=True)
        thread.start()

    def _generate(self):
        try:
            # Step 1: Extract text from all 3 PDFs
            self._log("Pas 1/3: Se extrage textul din PDF-uri...")

            self._log(f"  Anunț: {os.path.basename(self.notice_path.get())}")
            notice_pages = extract_pdf_text(self.notice_path.get())
            self._log(f"    {len(notice_pages)} pagini, {sum(len(p) for p in notice_pages)} caractere")

            self._log(f"  Fișa de date: {os.path.basename(self.datasheet_path.get())}")
            datasheet_pages = extract_pdf_text(self.datasheet_path.get())
            self._log(f"    {len(datasheet_pages)} pagini, {sum(len(p) for p in datasheet_pages)} caractere")

            self._log(f"  ATR: {os.path.basename(self.atr_path.get())}")
            atr_pages = extract_pdf_text(self.atr_path.get())
            self._log(f"    {len(atr_pages)} pagini, {sum(len(p) for p in atr_pages)} caractere")

            # Step 2: Generate Rezumat via Claude
            self._log("Pas 2/3: Se generează Rezumatul prin Claude API...")
            self._log(f"  Model: {self.app.model.get()}")

            company_data = {
                'leader': self.app.company_leader.get(),
                'associate': self.app.company_associate.get(),
                'subcontractor': self.app.company_subcontractor.get(),
                'warranty_months': self.app.warranty_months.get(),
                'pm_experience': self.app.pm_experience.get(),
            }

            rezumat_text = generate_rezumat(
                notice_pages, datasheet_pages, atr_pages, company_data,
                api_key=self.app.api_key.get(),
                model=self.app.model.get(),
                progress_callback=self._log
            )

            # Save raw text for reference
            raw_path = self.output_path.get().replace('.docx', '_raw.txt')
            with open(raw_path, 'w', encoding='utf-8') as f:
                f.write(rezumat_text)
            self._log(f"  Text brut salvat: {raw_path}")

            # Step 3: Build DOCX
            self._log("Pas 3/3: Se construiește documentul DOCX...")
            output_path = self.output_path.get()
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            build_docx(rezumat_text, output_path, doc_type="generic")
            self._log(f"  Document salvat: {output_path}")

            self._log("\nGata! Rezumatul a fost generat cu succes.")
            self.app.root.after(0, lambda: messagebox.showinfo(
                "Succes", f"Rezumat generat cu succes!\n\n{output_path}"))

        except anthropic.AuthenticationError:
            self._log("EROARE: Cheie API invalidă!")
            self.app.root.after(0, lambda: messagebox.showerror(
                "Eroare API", "Cheia API Anthropic este invalidă."))
        except anthropic.BadRequestError as e:
            self._log(f"EROARE API: {e}")
            self.app.root.after(0, lambda: messagebox.showerror("Eroare API", str(e)))
        except Exception as e:
            self._log(f"EROARE: {e}")
            self.app.root.after(0, lambda: messagebox.showerror("Eroare", str(e)))
        finally:
            self.generating = False
            self.app.root.after(0, lambda: self.gen_btn.config(state=tk.NORMAL))
            self.app.root.after(0, self.progress.stop)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def run():
    root = tk.Tk()
    app = App(root)
    root.mainloop()
