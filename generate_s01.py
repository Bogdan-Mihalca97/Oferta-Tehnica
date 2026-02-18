"""
Generate S01: Rezumat (Summary) - First generation unit
Reads extracted text from input PDFs and calls Claude API to generate
the Summary section of the Technical Execution Proposal.
"""
import json
import os
import anthropic
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load API key from .env file
def load_env(path):
    with open(path, 'r') as f:
        for line in f:
            line = line.strip()
            if '=' in line and not line.startswith('#'):
                key, val = line.split('=', 1)
                val = val.strip().strip('"').strip("'")
                os.environ[key.strip()] = val

BASE_DIR = r'c:\Users\mihal\Documents\testAchiztii'
load_env(os.path.join(BASE_DIR, '.env'))

# Load extracted context
with open(os.path.join(BASE_DIR, 's01_context.json'), 'r', encoding='utf-8') as f:
    context = json.load(f)

# Company data (configurable)
COMPANY_DATA = {
    "leader": "CRC AG S.R.L.",
    "associate": "CRC NEW ENERGY S.R.L.",
    "subcontractor": "BACKUP TECHNOLOGY S.R.L.",
    "warranty_months": 120,
    "pm_experience_projects": 5,
}

SYSTEM_PROMPT = """Ești un inginer expert român specializat în construcții de parcuri fotovoltaice și în elaborarea propunerilor tehnice pentru achiziții publice (SEAP/e-Licitație).

Scrii în limba română tehnică formală. Trebuie să:
- Folosești diacritice românești corect (ă, â, î, ș, ț)
- Referențiezi legislația română relevantă
- Urmezi structura impusă de Fișa de Date a achiziției
- Incluzi detalii tehnice specifice din documentația de proiect furnizată
- Folosești numele companiei și datele personalului furnizate
- Nu inventezi specificații tehnice - folosește doar datele din documentele de intrare

Formatul de ieșire:
- Folosește HEADING_1: pentru titluri de secțiuni principale
- Folosește HEADING_2: pentru subtitluri
- Folosește HEADING_3: pentru sub-subtitluri
- Folosește PARAGRAPH: pentru paragrafe normale
- Folosește BULLET: pentru elemente de listă
- Folosește BOLD_START: ... BOLD_END: pentru text bold inline
- Folosește TABLE: ... END_TABLE pentru tabele (cu | ca separator de coloane)
"""

USER_PROMPT = f"""Generează secțiunea 1: "Rezumat" (Summary) din Propunerea Tehnică de Execuție pentru un parc fotovoltaic.

Aceasta este prima secțiune a documentului și trebuie să conțină:

1.1. Date generale ale ofertantului
- Numele asocierii/companiei lider, asociat, subcontractant
- Structura asocierii și rolurile fiecărui membru

1.2. Obiectul contractului
- Descrierea succintă a obiectului achiziției (din Anunțul de participare)
- Valoarea estimată, coduri CPV
- Locația și beneficiarul

1.3. Avantaje competitive ale ofertei
- Perioada de garanție oferită: {COMPANY_DATA['warranty_months']} luni
- Experiența managerului de proiect: {COMPANY_DATA['pm_experience_projects']}+ proiecte similare
- Alte avantaje relevante (resurse proprii, certificări, experiență)

1.4. Abordarea generală a contractului
- Prezentare succintă a metodologiei de execuție
- Organizarea echipei
- Respectarea cerințelor din Caietul de sarcini și Fișa de date

DATELE COMPANIEI:
- Lider asociere: {COMPANY_DATA['leader']}
- Asociat: {COMPANY_DATA['associate']}
- Subcontractant: {COMPANY_DATA['subcontractor']}
- Garanție: {COMPANY_DATA['warranty_months']} luni
- Experiență MP: {COMPANY_DATA['pm_experience_projects']}+ proiecte

DOCUMENTE DE INTRARE:

=== ANUNȚ DE PARTICIPARE ===
{context['notice']}

=== FIȘA DE DATE ===
{context['datasheet']}

=== AVIZ TEHNIC DE RACORDARE (ATR) ===
{context['atr']}

EXEMPLU DE STIL (din documentul de referință - folosește acest stil dar cu datele din proiectul curent):
{context['reference_style']}

Generează aproximativ 3-4 pagini de conținut. Scrie în limba română, formal și tehnic.
"""

def parse_ai_output_to_docx(text, doc):
    """Parse the semi-structured AI output into DOCX elements."""
    lines = text.split('\n')

    for line in lines:
        line = line.rstrip()
        if not line:
            continue

        if line.startswith('HEADING_1:'):
            content = line[len('HEADING_1:'):].strip()
            p = doc.add_heading(content, level=1)
            p.style.font.name = 'Arial Narrow'
        elif line.startswith('HEADING_2:'):
            content = line[len('HEADING_2:'):].strip()
            p = doc.add_heading(content, level=2)
            p.style.font.name = 'Arial Narrow'
        elif line.startswith('HEADING_3:'):
            content = line[len('HEADING_3:'):].strip()
            p = doc.add_heading(content, level=3)
            p.style.font.name = 'Arial Narrow'
        elif line.startswith('PARAGRAPH:'):
            content = line[len('PARAGRAPH:'):].strip()
            content = _apply_bold(content, doc.add_paragraph())
        elif line.startswith('BULLET:'):
            content = line[len('BULLET:'):].strip()
            p = doc.add_paragraph(style='List Bullet')
            _apply_bold_to_paragraph(content, p)
        elif line.startswith('TABLE:'):
            # Collect table lines
            pass  # Tables handled in a second pass
        elif line.startswith('|'):
            # Table row - handled separately
            pass
        else:
            # Plain text - treat as paragraph
            if line.strip():
                doc.add_paragraph(line.strip())


def _apply_bold_to_paragraph(text, paragraph):
    """Handle BOLD_START/BOLD_END markers within text."""
    parts = text.split('BOLD_START:')
    if len(parts) == 1:
        paragraph.add_run(text)
        return

    # First part is normal
    if parts[0].strip():
        paragraph.add_run(parts[0])

    for part in parts[1:]:
        if 'BOLD_END:' in part:
            bold_text, rest = part.split('BOLD_END:', 1)
            run = paragraph.add_run(bold_text.strip())
            run.bold = True
            if rest.strip():
                paragraph.add_run(rest)
        else:
            run = paragraph.add_run(part.strip())
            run.bold = True


def _apply_bold(text, paragraph):
    """Apply bold formatting to a paragraph."""
    _apply_bold_to_paragraph(text, paragraph)
    return paragraph


def build_docx(ai_text, output_path):
    """Build a DOCX document from AI-generated text."""
    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # Parse AI output
    # Handle tables separately
    lines = ai_text.split('\n')
    table_buffer = []
    in_table = False

    for line in lines:
        line = line.rstrip()

        if line.strip() == 'TABLE:':
            in_table = True
            table_buffer = []
            continue
        elif line.strip() == 'END_TABLE':
            in_table = False
            if table_buffer:
                _add_table(doc, table_buffer)
            table_buffer = []
            continue

        if in_table:
            if line.strip().startswith('|'):
                table_buffer.append(line.strip())
            continue

        if not line.strip():
            continue

        if line.startswith('HEADING_1:'):
            content = line[len('HEADING_1:'):].strip()
            doc.add_heading(content, level=1)
        elif line.startswith('HEADING_2:'):
            content = line[len('HEADING_2:'):].strip()
            doc.add_heading(content, level=2)
        elif line.startswith('HEADING_3:'):
            content = line[len('HEADING_3:'):].strip()
            doc.add_heading(content, level=3)
        elif line.startswith('PARAGRAPH:'):
            content = line[len('PARAGRAPH:'):].strip()
            p = doc.add_paragraph()
            _apply_bold_to_paragraph(content, p)
        elif line.startswith('BULLET:'):
            content = line[len('BULLET:'):].strip()
            p = doc.add_paragraph(style='List Bullet')
            _apply_bold_to_paragraph(content, p)
        else:
            # Treat as plain paragraph
            if line.strip():
                doc.add_paragraph(line.strip())

    doc.save(output_path)
    print(f"Document saved to: {output_path}")


def _add_table(doc, rows):
    """Add a table to the document from pipe-separated rows."""
    parsed_rows = []
    for row in rows:
        cells = [c.strip() for c in row.strip('|').split('|')]
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    # Skip separator rows (like |---|---|)
    data_rows = [r for r in parsed_rows if not all(set(c) <= {'-', ' ', ':'} for c in r)]

    if not data_rows:
        return

    num_cols = max(len(r) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                table.cell(i, j).text = cell_text


def main():
    print("=" * 60)
    print("Generating S01: Rezumat (Summary)")
    print("=" * 60)

    # Initialize Claude client
    client = anthropic.Anthropic()

    print(f"Input tokens estimate: ~{len(USER_PROMPT) // 3} tokens")
    print("Calling Claude API...")

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8192,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": USER_PROMPT}],
    )

    ai_text = response.content[0].text

    # Save raw AI output for debugging
    raw_path = os.path.join(BASE_DIR, 'output', 's01_raw.txt')
    with open(raw_path, 'w', encoding='utf-8') as f:
        f.write(ai_text)
    print(f"Raw AI output saved to: {raw_path}")
    print(f"Output tokens: {response.usage.output_tokens}")
    print(f"Input tokens: {response.usage.input_tokens}")

    # Build DOCX
    docx_path = os.path.join(BASE_DIR, 'output', 'S01_Rezumat.docx')
    build_docx(ai_text, docx_path)

    print("\nDone!")


if __name__ == '__main__':
    main()
